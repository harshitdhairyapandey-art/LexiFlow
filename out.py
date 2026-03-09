import time
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PATHS ---
BASE_DIR = Path("LexiFlow")
FINAL_DIR = BASE_DIR / "5_final_novel"

# --- CONFIG ---
DEFAULT_FONT = 'Nirmala UI'
DEFAULT_SIZE = Pt(12)

# --- COMPILED REGEX ---
AI_PREFIX_RE = re.compile(r'^[A-Za-z\s]+:\s*', re.MULTILINE)

class OutputAssembler:
    def __init__(self, title="बियोंडार्स", author="LexiFlow", output_format="docx"):
        self.title = title
        self.author = author
        self.output_format = output_format
        self.doc = Document() if output_format == "docx" else None
        self.stats = {"files_merged": 0, "missing": 0}
        self._configure_styles()

    def _configure_styles(self):
        if not self.doc:
            return
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(0.75)

        style = self.doc.styles['Normal']
        f = style.font  # type: ignore
        f.name = DEFAULT_FONT
        f.size = DEFAULT_SIZE

        pf = style.paragraph_format  # type: ignore
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.15
        pf.space_after = Pt(10)
        pf.first_line_indent = Inches(0.3)

    def create_cover(self):
        if not self.doc:
            return
        for _ in range(5):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.title}\n")
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = DEFAULT_FONT

        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run("अनुवादित: आधुनिक हिंदी संस्करण\n")
        run_sub.font.size = Pt(16)
        run_sub.font.name = DEFAULT_FONT

        p_sub.add_run(f"दिनांक: {time.strftime('%Y-%m-%d')}\n")
        self.doc.add_page_break()

    def clean_content(self, text):
        """Filters out AI prefixes to keep the novel text clean."""
        return AI_PREFIX_RE.sub('', text).strip()

    def merge_files(self, folder_path, selected_files=None):
        """Main merge logic for stitching chunks into a Word doc."""
        folder = Path(folder_path)
        if not folder.exists():
            return

        all_files = sorted(folder.glob("y*.txt"))
        files_to_merge = all_files if not selected_files else [
            folder / f for f in selected_files if (folder / f).exists()
        ]

        if not files_to_merge:
            return

        if self.doc:
            self.create_cover()

        for i, file in enumerate(files_to_merge, 1):
            try:
                content = file.read_text(encoding="utf-8", errors="ignore")

                if self.doc:
                    chapter_num = int(file.stem.replace('y', ''))
                    h = self.doc.add_heading(f"अध्याय {chapter_num}", level=2)
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for para in re.split(r"\n\s*\n", content):
                        cleaned = self.clean_content(para)
                        if cleaned:
                            self.doc.add_paragraph(cleaned)

                    if i != len(files_to_merge):
                        self.doc.add_page_break()

                self.stats["files_merged"] += 1
            except Exception as e:
                print(f"Skipped {file.name}: {e}")

        self._save_output()

    def _save_output(self):
        FINAL_DIR.mkdir(parents=True, exist_ok=True)
        timestamp = time.strftime("%Y%m%d_%H%M")
        if self.doc:
            filename = f"Final_Novel_{timestamp}.docx"
            save_path = FINAL_DIR / filename
            self.doc.save(str(save_path))
            print(f"✅ Exported: {filename}")

# --- MANUAL RUN ---
if __name__ == "__main__":
    assembler = OutputAssembler()
    source = BASE_DIR / "4_polished_output"
    if not any(source.glob("y*.txt")):
        source = BASE_DIR / "3_translated_chunks"
    assembler.merge_files(source)